"""Turn an uploaded resume file (PDF / DOCX / TXT) into plain text."""

import io
import re

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class ExtractionError(ValueError):
    """Raised when a file cannot be parsed into usable text."""


def _clean(text: str) -> str:
    # Normalise the whitespace soup that PDF extraction usually produces.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # corrupt / encrypted file
        raise ExtractionError(f"Could not read PDF: {exc}") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError("PDF is password protected.") from exc

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    # Resumes often keep skills/dates inside tables, which paragraphs miss.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode text file.")


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file extension and return cleaned plain text."""
    if not data:
        raise ExtractionError("Uploaded file is empty.")

    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _from_pdf(data)
    elif name.endswith(".docx"):
        text = _from_docx(data)
    elif name.endswith((".txt", ".md")):
        text = _from_txt(data)
    elif name.endswith(".doc"):
        raise ExtractionError("Legacy .doc is not supported — please save as .docx or PDF.")
    else:
        raise ExtractionError(
            f"Unsupported file type. Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    cleaned = _clean(text)
    if len(cleaned) < 100:
        raise ExtractionError(
            "Extracted almost no text. If this is a scanned/image PDF, "
            "export a text-based PDF or paste the resume text instead."
        )
    return cleaned
